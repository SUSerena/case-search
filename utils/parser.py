"""
Excel/Word项目信息表解析器
解析"项目信息表"格式的Excel文件，提取结构化数据
"""
import pandas as pd
import re
import os
from utils.db import insert_project, add_drawing, add_photo, add_materials
from utils.location import detect_location


def parse_project_excel(file_path):
    """
    解析项目信息表 Excel 文件
    返回项目数据字典
    """
    try:
        df = pd.read_excel(file_path, sheet_name=0, header=None)
    except Exception as e:
        raise Exception(f"读取Excel失败: {e}")

    data = {
        'project_name': '',
        'project_type': '',
        'sales_person': '',
        'confirm_date': '',
        'location': '',
        'coordinates': '',
        'description': '',
        'sea_conditions': {},
        'cage_params': {},
        'platform_params': {},
        'breakwater_params': {}
    }

    # 遍历每一行，根据字段标签提取值
    for i in range(len(df)):
        row = df.iloc[i].tolist()
        # 转换为字符串列表
        cells = [str(x).strip() if pd.notna(x) else '' for x in row]

        # 项目名称
        if '项目名称' in cells[0]:
            # 找后续列中的值
            for j in range(1, len(cells)):
                if cells[j] and cells[j] != 'NaT':
                    data['project_name'] = cells[j]
                    break

        # 业务人员
        elif '业务人员' in cells[0]:
            for j in range(1, len(cells)):
                if cells[j] and cells[j] != 'NaT':
                    data['sales_person'] = cells[j]
                    break

        # 项目类型
        elif '项目类型' in cells[0]:
            for j in range(1, len(cells)):
                if cells[j] and cells[j] != 'NaT':
                    data['project_type'] = cells[j]
                    break

        # 海域坐标
        elif '海域坐标' in cells[0]:
            for j in range(1, len(cells)):
                if cells[j] and cells[j] != 'NaT' and '陆地施工' not in cells[j]:
                    data['coordinates'] = cells[j]
                    break

        # 满潮水深
        elif '满潮水深' in cells[0]:
            if len(cells) > 1 and cells[1] and cells[1] != 'NaT':
                data['sea_conditions']['water_depth'] = _extract_max_depth(cells[1])
            # 流向
            if len(cells) > 3 and cells[3] and '流向' in str(cells[3]):
                if len(cells) > 4 and cells[4]:
                    data['sea_conditions']['flow_direction'] = cells[4]

        # 水位差
        elif '水位差' in cells[0]:
            if len(cells) > 1 and cells[1] and cells[1] != 'NaT':
                data['sea_conditions']['water_level_diff'] = _extract_number(cells[1])
            # 底质
            if len(cells) > 3 and cells[3] and '底质' in str(cells[3]):
                if len(cells) > 4 and cells[4]:
                    data['sea_conditions']['seabed_type'] = cells[4]

        # 最大波高
        elif '最大波高' in cells[0]:
            if len(cells) > 1 and cells[1] and cells[1] != 'NaT':
                data['sea_conditions']['max_wave_height'] = _extract_number(cells[1])
            # 最大流速
            if len(cells) > 3 and cells[3] and '最大流速' in str(cells[3]):
                if len(cells) > 4 and cells[4]:
                    data['sea_conditions']['max_flow_speed'] = _extract_number(cells[4])

        # 常风向
        elif '常风向' in cells[0]:
            if len(cells) > 1 and cells[1] and cells[1] != 'NaT':
                data['sea_conditions']['common_wind_direction'] = cells[1]
            # 最大风向
            if len(cells) > 3 and cells[3] and '最大风向' in str(cells[3]):
                if len(cells) > 4 and cells[4]:
                    data['sea_conditions']['max_wind_direction'] = cells[4]

        # 最大风速（单独一行的情况）
        elif '最大风速' in cells[0] and len(cells[0]) < 10:
            if len(cells) > 1 and cells[1] and cells[1] != 'NaT':
                data['sea_conditions']['max_wind_speed'] = _extract_number(cells[1])
        elif len(cells) > 3 and cells[3] and '最大风速' in str(cells[3]):
            if len(cells) > 4 and cells[4] and cells[4] != 'NaT':
                data['sea_conditions']['max_wind_speed'] = _extract_number(cells[4])

        # 网箱管径
        # 需要先识别哪一行是网箱类管径行
        pass

    # 专门解析网箱/平台/防波堤参数区域（从"项目信息"之后的行开始）
    in_project_info = False
    for i in range(len(df)):
        row = df.iloc[i].tolist()
        cells = [str(x).strip() if pd.notna(x) else '' for x in row]

        if '项目信息' in cells[0]:
            in_project_info = True
            continue

        if not in_project_info:
            continue

        # 到达确认时间行则停止
        if '确认时间' in ''.join(cells):
            # 提取确认时间
            for j in range(len(cells)):
                if cells[j] and '确认时间' not in cells[j] and cells[j] != 'NaT':
                    try:
                        data['confirm_date'] = cells[j].split()[0] if ' ' in cells[j] else cells[j]
                    except:
                        pass
            break

        # 管径行（列0=网箱标签, 列1=网箱值; 列2=平台标签, 列3=平台值; 列4=防波堤标签, 列5=防波堤值）
        if cells[0] == '管径*':
            # 网箱管径（列1）
            if len(cells) > 1 and cells[1] and cells[1] != 'NaT':
                data['cage_params']['pipe_diameter'] = cells[1]
            # 平台管径（列3）
            if len(cells) > 3 and cells[3] and cells[3] != 'NaT':
                data['platform_params']['pipe_diameter'] = cells[3]
            # 防波堤管径（列5）
            if len(cells) > 5 and cells[5] and cells[5] != 'NaT':
                data['breakwater_params']['pipe_diameter'] = cells[5]

        # 周长/直径
        elif '周长' in cells[0] and '直径' in cells[0]:
            # 网箱周长/直径（列1）
            if len(cells) > 1 and cells[1] and cells[1] != 'NaT':
                data['cage_params']['_perimeter_text'] = cells[1]
                perim = _extract_perimeter(cells[1])
                if perim:
                    data['cage_params']['perimeter'] = perim
                dia = _extract_diameter(cells[1])
                if dia:
                    data['cage_params']['diameter'] = dia
                cnt = _extract_cage_count(cells[1])
                if cnt:
                    data['cage_params']['cage_count'] = cnt
                ctype = _extract_cage_type(cells[1])
                if ctype:
                    data['cage_params']['cage_type'] = ctype
            # 平台尺寸规格（列3）
            if len(cells) > 3 and cells[3] and cells[3] != 'NaT':
                data['platform_params']['size_spec'] = cells[3]
            # 防波堤规格选型（列5）
            if len(cells) > 5 and cells[5] and cells[5] != 'NaT':
                data['breakwater_params']['spec_type'] = cells[5]

        # 支架间距
        elif cells[0] == '支架间距':
            # 网箱（列1）
            if len(cells) > 1 and cells[1] and cells[1] != 'NaT':
                data['cage_params']['bracket_spacing'] = cells[1]
            # 防波堤（列5）
            if len(cells) > 5 and cells[5] and cells[5] != 'NaT':
                data['breakwater_params']['bracket_spacing'] = cells[5]

        # 过道宽度
        elif '过道宽度' in cells[0]:
            if len(cells) > 1 and cells[1] and cells[1] != 'NaT':
                data['cage_params']['walkway_width'] = cells[1]

        # 系缆加强套管
        elif '系缆加强套管' in cells[0]:
            if len(cells) > 1 and cells[1] and cells[1] != 'NaT':
                data['cage_params']['mooring_sleeve'] = cells[1]

        # 锚点数量
        elif '锚点数量' in cells[0]:
            if len(cells) > 1 and cells[1] and cells[1] != 'NaT':
                data['cage_params']['anchor_point_count'] = _extract_number(cells[1])

        # 网衣需求
        elif '网衣需求' in cells[0]:
            if len(cells) > 1 and cells[1] and cells[1] != 'NaT':
                data['cage_params']['net_demand'] = cells[1]

        # 造型要求（平台类，列2是标签，列3是值）
        elif len(cells) > 2 and cells[2] == '造型要求':
            if len(cells) > 3 and cells[3] and cells[3] != 'NaT':
                data['platform_params']['shape_requirement'] = cells[3]

        # 附属设施（平台类，列2是标签，列3是值）
        elif len(cells) > 2 and '附属设施' in str(cells[2]):
            if len(cells) > 3 and cells[3] and cells[3] != 'NaT':
                data['platform_params']['附属设施'] = cells[3]

        # 承载力要求（平台类，列2是标签，列3是值）
        elif len(cells) > 2 and '承载力' in str(cells[2]):
            if len(cells) > 3 and cells[3] and cells[3] != 'NaT':
                data['platform_params']['bearing_capacity'] = cells[3]

        # 靠泊船只（平台类，列2是标签，列3是值）
        elif len(cells) > 2 and '靠泊船只' in str(cells[2]):
            if len(cells) > 3 and cells[3] and cells[3] != 'NaT':
                data['platform_params']['docking_ships'] = cells[3]

        # 减波率要求（防波堤类，列4是标签，列5是值）
        elif len(cells) > 4 and '减波率' in str(cells[4]):
            if len(cells) > 5 and cells[5] and cells[5] != 'NaT':
                data['breakwater_params']['wave_reduction_rate'] = cells[5]

        # 其他特殊要求 - 合并到描述
        elif '其他特殊要求' in cells[0]:
            descs = []
            for j in range(1, len(cells)):
                if cells[j] and cells[j] != 'NaT' and '其他特殊要求' not in cells[j]:
                    descs.append(cells[j])
            if descs:
                data['description'] = '\n'.join(descs)

    # 从项目名称中提取地点信息
    loc = _extract_location(data['project_name'])
    if loc:
        data['location'] = loc

    # 智能识别地域信息（城市、省份、海域）
    text_for_detect = data['project_name'] + ' ' + (data.get('description') or '') + ' ' + (data.get('location') or '')
    city, province, sea_area = detect_location(text_for_detect)
    if city:
        # 如果已有location但更精确，保留原location；否则用识别到的city
        if not data.get('location') or len(data['location']) < 2:
            data['location'] = city
    if province:
        data['province'] = province
    if sea_area:
        data['sea_area'] = sea_area

    return data


def parse_materials_excel(file_path):
    """
    解析材料统计表
    返回材料列表
    """
    try:
        df = pd.read_excel(file_path, sheet_name=0, header=None)
    except Exception as e:
        raise Exception(f"读取材料统计Excel失败: {e}")

    materials = []
    current_category = ''

    # 找到表头行（包含"货品编码"或"材料名称"的行）
    header_row = -1
    for i in range(min(10, len(df))):
        row_str = ' '.join([str(x) for x in df.iloc[i].tolist()])
        if '材料名称' in row_str or '货品编码' in row_str:
            header_row = i
            break

    if header_row < 0:
        return materials

    # 从表头下一行开始读取
    for i in range(header_row + 1, len(df)):
        row = df.iloc[i].tolist()
        cells = [x if pd.notna(x) else '' for x in row]

        # 检测分类行（只有第一列有值，且不是货品编码）
        first_col = str(cells[0]).strip() if cells[0] else ''
        has_content = any(str(c).strip() for c in cells[1:] if str(c).strip() and str(c) != 'nan')

        # 如果只有第一列有内容，可能是分类标题
        if first_col and not has_content and len(first_col) < 20:
            # 检查是不是纯数字编码（货品编码）
            if not re.match(r'^\d+$', first_col):
                current_category = first_col
                continue

        # 正常数据行：有材料名称
        material_name = ''
        if len(cells) > 1:
            material_name = str(cells[1]).strip() if cells[1] else ''

        if not material_name:
            continue

        material = {
            'category': current_category,
            'material_name': material_name,
            'model': str(cells[2]).strip() if len(cells) > 2 and cells[2] else '',
            'unit': str(cells[3]).strip() if len(cells) > 3 and cells[3] else '',
            'quantity': None,
            'unit_weight': None,
            'total_weight': None
        }

        # 数量（总计列，通常是第8列 index=8）
        if len(cells) > 8 and cells[8]:
            try:
                material['quantity'] = float(cells[8])
            except (ValueError, TypeError):
                pass

        # 单重
        if len(cells) > 9 and cells[9]:
            try:
                material['unit_weight'] = float(cells[9])
            except (ValueError, TypeError):
                pass

        # 总重
        if len(cells) > 10 and cells[10]:
            try:
                material['total_weight'] = float(cells[10])
            except (ValueError, TypeError):
                pass

        materials.append(material)

    return materials


def _extract_number(text):
    """从文本中提取第一个数值"""
    if not text:
        return None
    text = str(text)
    match = re.search(r'(\d+\.?\d*)', text)
    if match:
        return float(match.group(1))
    return None


def _extract_max_depth(text):
    """从水深描述中提取最大水深值，如'最低潮5m，满潮7-8m' -> 8.0"""
    if not text:
        return None
    text = str(text)
    # 提取所有数字
    numbers = re.findall(r'(\d+\.?\d*)', text)
    if not numbers:
        return None
    # 返回最大值
    return max(float(n) for n in numbers)


def _extract_perimeter(text):
    """提取周长数值，如'90m周长'、'周长90m'"""
    if not text:
        return None
    text = str(text)
    # 匹配 xxx m周长 or 周长xxx m or xxxm周长
    patterns = [
        r'(\d+\.?\d*)\s*[mM米]\s*周长',
        r'周长\s*(\d+\.?\d*)\s*[mM米]?',
        r'C\s*(\d+\.?\d*)',
    ]
    for p in patterns:
        match = re.search(p, text)
        if match:
            return float(match.group(1))
    return None


def _extract_diameter(text):
    """提取直径数值"""
    if not text:
        return None
    text = str(text)
    patterns = [
        r'直径\s*(\d+\.?\d*)\s*[mM米]?',
        r'(\d+\.?\d*)\s*[mM米]\s*直径',
    ]
    for p in patterns:
        match = re.search(p, text)
        if match:
            return float(match.group(1))
    return None


def _extract_cage_count(text):
    """提取网箱数量，如'3口'、'18口'"""
    if not text:
        return None
    text = str(text)
    match = re.search(r'(\d+)\s*口', text)
    if match:
        return int(match.group(1))
    # "3个"
    match = re.search(r'(\d+)\s*个', text)
    if match:
        return int(match.group(1))
    return None


def _extract_cage_type(text):
    """提取网箱类型，如'圆形'、'方形'、'韩式'"""
    if not text:
        return ''
    text = str(text)
    types = []
    if '圆形' in text:
        types.append('圆形')
    if '方形' in text:
        types.append('方形')
    if '韩式' in text:
        types.append('韩式')
    if '九宫格' in text:
        types.append('九宫格')
    return '、'.join(types) if types else ''


def _extract_location(project_name):
    """从项目名称中提取地点"""
    if not project_name:
        return ''
    # 常见地名模式
    # 提取"连云港"、"青岛"、"大连"等地名
    # 简单方法：取前2-4个字符中包含地名的部分
    locations = ['连云港', '青岛', '大连', '烟台', '威海', '日照', '宁波', '温州',
                 '福州', '厦门', '深圳', '珠海', '湛江', '北海', '三亚', '舟山',
                 '台州', '泉州', '汕头', '阳江', '茂名', '海口', '儋州', '文昌',
                 '秦皇岛', '葫芦岛', '丹东', '营口', '盘锦', '锦州', '朝阳',
                 '上海', '天津', '广州', '杭州', '南京', '苏州']
    for loc in locations:
        if loc in project_name:
            return loc
    # 默认提取前3个字符作为地点候选
    if len(project_name) >= 3:
        return project_name[:3]
    return project_name


def import_project_from_excel(excel_path, drawings_dir=None, photos_dir=None):
    """
    从Excel文件导入项目，自动关联图纸和照片
    返回: 项目ID
    """
    # 解析项目信息
    project_data = parse_project_excel(excel_path)

    if not project_data['project_name']:
        raise Exception("未能识别出项目名称")

    # 检查是否已存在
    from utils.db import project_exists
    if project_exists(project_data['project_name']):
        raise Exception(f"项目 '{project_data['project_name']}' 已存在")

    # 插入项目
    project_id = insert_project(project_data)

    # 关联图纸
    if drawings_dir and os.path.isdir(drawings_dir):
        project_keywords = project_data['project_name'][:4]  # 取项目名前几个字作为关键词
        for f in os.listdir(drawings_dir):
            fpath = os.path.join(drawings_dir, f)
            if os.path.isfile(fpath):
                # 判断是否相关（文件名包含项目关键词）
                fname_lower = f.lower()
                if any(kw in f for kw in [project_data['project_name'][:4], project_data['location'][:2]] if kw):
                    ftype = 'pdf' if f.lower().endswith('.pdf') else \
                            ('image' if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')) else 'other')
                    fsize = os.path.getsize(fpath)
                    # 识别图纸类型
                    dtype = ''
                    if '框架' in f:
                        dtype = '框架图'
                    elif '锚固' in f:
                        dtype = '锚固图'
                    elif '平面' in f or '布局' in f:
                        dtype = '总平面图'
                    elif '材料' in f:
                        dtype = '材料统计'

                    add_drawing(project_id, f, fpath, ftype, fsize, dtype)

    # 关联照片
    if photos_dir and os.path.isdir(photos_dir):
        for f in os.listdir(photos_dir):
            fpath = os.path.join(photos_dir, f)
            if os.path.isfile(fpath) and f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                fsize = os.path.getsize(fpath)
                add_photo(project_id, f, fpath, 'image', fsize, '项目照片')

    return project_id


# ====================== 智能识别导入 ======================

# 字段关键词映射表：关键词 → 字段路径
# 每条：(关键词列表, 数据路径)
FIELD_KEYWORDS = [
    # 基本信息
    (['项目名称', '工程名称', '项目名'], ['project_name']),
    (['业务人员', '负责人', '业务员', '业务代表', '销售'], ['sales_person']),
    (['项目类型', '产品类型', '类型'], ['project_type']),
    (['确认时间', '签订日期', '合同日期'], ['confirm_date']),
    (['项目地点', '位置', '所在地', '海域', '地点', '施工地点', '工程地点'], ['location']),
    (['坐标', '经纬度'], ['coordinates']),
    (['工期', '施工工期', '合同工期'], ['construction_period']),
    (['预算', '造价', '投资', '合同金额', '中标价'], ['budget']),
    (['数量', '网箱数量', '台数', '套数'], ['quantity']),
    (['面积', '养殖面积', '海域面积'], ['area']),
    # 海况参数
    (['水深', '深度', '满潮水深', '设计水深'], ['sea_conditions', 'water_depth']),
    (['水位差', '潮差'], ['sea_conditions', 'water_level_diff']),
    (['波高', '波浪', '最大波高'], ['sea_conditions', 'max_wave_height']),
    (['流速', '海流', '最大流速'], ['sea_conditions', 'max_flow_speed']),
    (['流向'], ['sea_conditions', 'flow_direction']),
    (['风速', '最大风速'], ['sea_conditions', 'max_wind_speed']),
    (['风向', '常风向'], ['sea_conditions', 'common_wind_direction']),
    (['底质', '海底', '海床', '地质'], ['sea_conditions', 'seabed_type']),
    # 网箱参数
    (['管径', '主管径', 'DN'], ['cage_params', 'pipe_diameter']),
    (['周长', '直径', '周长直径'], ['cage_params', '_perimeter_text']),
    (['网箱类型', '网箱形式'], ['cage_params', 'cage_type']),
    (['网箱数量', '网箱数'], ['cage_params', 'cage_count']),
    (['支架间距'], ['cage_params', 'bracket_spacing']),
    (['过道宽度', '走道宽度'], ['cage_params', 'walkway_width']),
    (['系缆加强', '套管'], ['cage_params', 'mooring_sleeve']),
    (['锚点数量', '锚点数'], ['cage_params', 'anchor_point_count']),
    (['网衣需求', '网衣'], ['cage_params', 'net_demand']),
    # 平台参数
    (['造型要求', '平台造型'], ['platform_params', 'shape_requirement']),
    (['承载力', '承载能力'], ['platform_params', 'bearing_capacity']),
    (['靠泊船只', '靠泊'], ['platform_params', 'docking_ships']),
    (['附属设施'], ['platform_params', '附属设施']),
    # 防波堤参数
    (['减波率', '减波'], ['breakwater_params', 'wave_reduction_rate']),
    (['防波堤规格', '规格选型'], ['breakwater_params', 'spec_type']),
    # 描述
    (['其他要求', '特殊要求', '备注', '说明'], ['description']),
]


def parse_smart(file_path):
    """
    智能识别导入：自动识别Excel/Word/PDF/TXT中的项目信息
    不依赖固定模板格式，通过关键词模糊匹配提取字段
    返回: (data, raw_text, file_format)
    """
    ext = os.path.splitext(file_path)[1].lower()
    raw_text = ''
    file_format = ''

    if ext in ('.xlsx', '.xls'):
        file_format = 'excel'
        raw_text = _extract_excel_text(file_path)
    elif ext == '.docx':
        file_format = 'word'
        raw_text = _extract_word_text(file_path)
    elif ext == '.pdf':
        file_format = 'pdf'
        raw_text = _extract_pdf_text(file_path)
    elif ext == '.txt':
        file_format = 'txt'
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            raw_text = f.read()
    else:
        raise Exception(f"不支持的文件格式: {ext}")

    if not raw_text.strip():
        raise Exception("文件内容为空，无法识别")

    # 智能识别字段
    data = _smart_match_fields(raw_text)

    # 补充地域信息
    text_for_detect = ' '.join([str(v) for v in [
        data.get('project_name', ''), data.get('location', ''), data.get('description', '')
    ] if v])
    city, province, sea_area = detect_location(text_for_detect)
    if city and (not data.get('location') or len(data['location']) < 2):
        data['location'] = city
    if province:
        data['province'] = province
    if sea_area:
        data['sea_area'] = sea_area

    return data, raw_text, file_format


def _extract_excel_text(file_path):
    """提取Excel所有sheet的文本内容"""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, data_only=True)
        lines = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines.append(f'[{sheet_name}]')
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() if c is not None else '' for c in row]
                if any(c for c in cells):
                    lines.append('\t'.join(cells))
        return '\n'.join(lines)
    except Exception as e:
        raise Exception(f"读取Excel失败: {e}（可能需要安装 openpyxl: pip install openpyxl）")


def _extract_word_text(file_path):
    """提取Word文档文本"""
    try:
        from docx import Document
        doc = Document(file_path)
        lines = []
        # 段落
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        # 表格
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(c for c in cells):
                    lines.append('\t'.join(cells))
        return '\n'.join(lines)
    except ImportError:
        raise Exception("解析Word需要安装 python-docx: pip install python-docx")
    except Exception as e:
        raise Exception(f"读取Word失败: {e}")


def _extract_pdf_text(file_path):
    """提取PDF文本"""
    try:
        import pdfplumber
        lines = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    lines.append(text)
                # 也提取表格
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        cells = [str(c).strip() if c else '' for c in row]
                        if any(c for c in cells):
                            lines.append('\t'.join(cells))
        return '\n'.join(lines)
    except ImportError:
        raise Exception("解析PDF需要安装 pdfplumber: pip install pdfplumber")
    except Exception as e:
        raise Exception(f"读取PDF失败: {e}")


def _smart_match_fields(text):
    """用关键词模糊匹配从文本中提取字段值"""
    lines = text.split('\n')
    data = {
        'project_name': '', 'project_type': '', 'sales_person': '',
        'confirm_date': '', 'location': '', 'coordinates': '',
        'description': '', 'construction_period': '', 'budget': '',
        'quantity': '', 'area': '',
        'sea_conditions': {}, 'cage_params': {},
        'platform_params': {}, 'breakwater_params': {}
    }

    matched_fields = set()  # 已匹配字段，避免重复

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 尝试用冒号/制表符分割为 key:value 对
        parts = _split_kv(line)
        if not parts:
            continue

        for key, value in parts:
            value = str(value).strip()
            if not value or value == 'None' or value == 'nan':
                continue

            for keywords, path in FIELD_KEYWORDS:
                field_key = '.'.join(path)
                if field_key in matched_fields:
                    continue

                # 检查key是否包含任何关键词
                if any(kw in key for kw in keywords):
                    _set_nested(data, path, _clean_value(value, path[-1]))
                    matched_fields.add(field_key)
                    break

    return data


def _split_kv(line):
    """
    将一行文本拆分为(key, value)对列表
    支持: "key: value" / "key\tvalue" / "key：value"
    也支持一行多对: "key1\tvalue1\tkey2\tvalue2"
    """
    pairs = []

    # 先尝试按冒号分割
    if '：' in line or ': ' in line:
        if '：' in line:
            parts = line.split('：', 1)
        else:
            parts = line.split(': ', 1)
        if len(parts) == 2:
            pairs.append((parts[0].strip(), parts[1].strip()))
            return pairs

    # 按制表符分割
    if '\t' in line:
        cells = [c.strip() for c in line.split('\t')]
        # 相邻两两配对: key, value, key, value...
        i = 0
        while i + 1 < len(cells):
            if cells[i] and cells[i + 1]:
                pairs.append((cells[i], cells[i + 1]))
            i += 2
        return pairs

    # 按连续空格分割
    cells = line.split()
    if len(cells) >= 2:
        i = 0
        while i + 1 < len(cells):
            if cells[i] and cells[i + 1]:
                pairs.append((cells[i], cells[i + 1]))
            i += 2

    return pairs


def _set_nested(data, path, value):
    """在嵌套dict中设置值"""
    if len(path) == 1:
        data[path[0]] = value
    else:
        key = path[0]
        if key not in data:
            data[key] = {}
        data[key][path[1]] = value


def _clean_value(value, field_name):
    """清理字段值，根据字段类型做适当处理"""
    if not value:
        return ''
    # 数值类字段：提取数字
    numeric_fields = ['water_depth', 'water_level_diff', 'max_wave_height',
                      'max_flow_speed', 'max_wind_speed', 'cage_count',
                      'anchor_point_count']
    if field_name in numeric_fields:
        num = _extract_number(value)
        if num is not None:
            return num
    # 管径字段：保留原文
    if field_name == 'pipe_diameter':
        # 去掉"DN"前缀保留数字
        m = re.search(r'DN?\s*(\d+)', str(value))
        if m:
            return f'DN{m.group(1)}'
        return str(value)
    # 周长文本：保留原文供后续提取
    if field_name == '_perimeter_text':
        text = str(value)
        perim = _extract_perimeter(text)
        if perim:
            return text  # 保留原文，后续在insert时处理
        return text
    return str(value).strip()

