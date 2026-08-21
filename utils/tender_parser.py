"""
招标文件智能分析模块
从招标文件（Word/PDF/TXT）中自动提取项目关键信息
增强版：支持范围值、前缀符号、多种表述方式
"""
import re
import os
from utils.location import detect_location


def parse_tender_document(file_path):
    """
    解析招标文件，提取关键信息
    返回: dict 项目信息
    """
    text = _read_document_text(file_path)
    if not text:
        return None

    result = {
        'project_name': '',
        'project_type': '',
        'location': '',
        'province': '',
        'sea_area': '',
        'project_scale': '',
        'construction_period': '',
        'budget': '',
        'tenderer': '',
        'description': '',
        'sea_conditions': {},
        'cage_params': {},
        'cage_params_circle': {},
        'cage_params_square': {},
        'platform_params': {},
        'breakwater_params': {},
        'goods_requirements': [],
        'summary': ''
    }

    # 1. 提取项目名称
    result['project_name'] = _extract_project_name(text, file_path)

    # 2. 识别项目类型
    result['project_type'] = _extract_project_type(text)

    # 3. 提取招标人
    result['tenderer'] = _extract_tenderer(text)

    # 4. 识别地域信息
    city, province, sea_area = detect_location(text)
    if city:
        result['location'] = city
    if province:
        result['province'] = province
    if sea_area:
        result['sea_area'] = sea_area

    # 5. 提取交货地点
    delivery_location = _extract_delivery_location(text)
    result['delivery_location'] = delivery_location
    if delivery_location and not result['location']:
        result['location'] = delivery_location

    # 6. 提取项目规模
    result['project_scale'] = _extract_project_scale(text)

    # 7. 提取工期
    result['construction_period'] = _extract_construction_period(text)

    # 8. 提取预算/投资金额
    result['budget'] = _extract_budget(text)

    # 9. 提取海况参数（增强）
    result['sea_conditions'] = _extract_sea_conditions(text)

    # 10. 提取网箱参数（增强，区分圆形和方形）
    result['cage_params'] = _extract_cage_params(text)
    # 根据 cage_shape 分配到 circle 或 square
    shape = result['cage_params'].get('cage_shape', 'circle')
    if shape == 'square':
        result['cage_params_square'] = result['cage_params']
    else:
        result['cage_params_circle'] = result['cage_params']

    # 11. 平台/防波堤参数
    result['platform_params'] = _extract_platform_params(text)
    result['breakwater_params'] = _extract_breakwater_params(text)

    # 12. 提取货物需求（重点）
    result['goods_requirements'] = _extract_goods_requirements(text)
    result['summary'] = _generate_goods_summary(result['goods_requirements'])

    # 13. 生成项目概述
    result['description'] = _generate_summary(text, result)

    return result


# ====================== 文本读取 ======================

def _read_document_text(file_path):
    """读取文档文本内容，支持txt/docx/pdf"""
    if not file_path or not os.path.exists(file_path):
        return ''

    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

        elif ext == '.docx':
            try:
                from docx import Document
                doc = Document(file_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if cells:
                            paragraphs.append('\t'.join(cells))
                return '\n'.join(paragraphs)
            except ImportError:
                raise Exception("解析Word需要安装 python-docx: pip install python-docx")

        elif ext == '.pdf':
            text = ''
            # 方法1: pdfplumber（优先，表格提取更好）
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + '\n'
                        # 提取表格
                        tables = page.extract_tables()
                        for table in tables:
                            for row in table:
                                cells = [str(c).strip() if c else '' for c in row]
                                if any(c for c in cells):
                                    text += '\t'.join(cells) + '\n'
                if text.strip():
                    return text
            except ImportError:
                pass

            # 方法2: PyPDF2
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + '\n'
                return text
            except ImportError:
                raise Exception("解析PDF需要安装 pdfplumber 或 PyPDF2: pip install pdfplumber")

        elif ext == '.doc':
            raise Exception(".doc旧格式暂不支持，请转换为.docx或.pdf")
        else:
            raise Exception(f"不支持的文件格式: {ext}")

    except Exception as e:
        raise Exception(f"文件读取失败: {e}")


# ====================== 项目信息提取 ======================

def _extract_project_name(text, file_path=None):
    """提取项目名称"""
    # 常见模式
    patterns = [
        r'项目名称[：:]\s*([^\n，。；;]{3,50})',
        r'工程名称[：:]\s*([^\n，。；;]{3,50})',
        r'招标项目[：:]\s*([^\n，。；;]{3,50})',
        r'([\u4e00-\u9fa5A-Za-z0-9（）()]{5,50}(?:网箱|平台|防波堤|渔场|养殖|海洋牧场)(?:项目|工程|采购))',
    ]

    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            name = match.group(1).strip()
            name = re.sub(r'[（(].*?[）)]', '', name).strip()
            if 3 <= len(name) <= 50:
                return name

    # 从文件名提取
    if file_path:
        fname = os.path.splitext(os.path.basename(file_path))[0]
        # 清理文件名中的hash前缀
        fname = re.sub(r'^[a-f0-9]+_', '', fname)
        if len(fname) > 5:
            return fname[:50]

    # 第一行
    first_line = text.strip().split('\n')[0][:50] if text.strip() else '未命名项目'
    return first_line


def _extract_project_type(text):
    """识别项目类型"""
    types = []
    if re.search(r'网箱|养殖|渔排', text):
        types.append('网箱')
    if re.search(r'平台|海洋牧场|综合体|渔旅|休闲', text):
        types.append('平台')
    if re.search(r'防波堤|消波|减波', text):
        types.append('防波堤')

    if not types:
        return '网箱'

    return '＋'.join(types)


def _extract_tenderer(text):
    """提取招标人"""
    patterns = [
        r'招\s*标\s*人[：:]\s*([^\n，。；;]{4,30})',
        r'采\s*购\s*人[：:]\s*([^\n，。；;]{4,30})',
        r'项目业主[：:]\s*([^\n，。；;]{4,30})',
        r'招\s*标\s*代\s*理[：:]\s*([^\n，。；;]{4,30})',
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1).strip()
    return ''


def _extract_delivery_location(text):
    """提取交货地点"""
    patterns = [
        r'交货地点[：:]\s*([^\n，。；;]{4,50})',
        r'交货地点[：:](.*?)(?:\n|；)',
        r'项目地点[：:]\s*([^\n，。；;]{4,50})',
        r'位于([^\n，。；;]{4,30}海域)',
        r'施工地点[：:]\s*([^\n，。；;]{4,50})',
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            loc = match.group(1).strip()
            if len(loc) <= 50:
                return loc
    return ''


def _extract_project_scale(text):
    """提取项目规模"""
    patterns = [
        r'(海域面积\d+(?:\.\d+)?(?:亩|平方米|㎡|公顷))',
        r'项目规模[：:]\s*([^\n，。；;]{5,80})',
        r'建设规模[：:]\s*([^\n，。；;]{5,80})',
        r'(计划投放\d+座[^。；;]+)',
        r'((?:投放|建设|安装)\d+座[^。；;\n]{5,60})',
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1).strip()[:80]
    return ''


def _extract_construction_period(text):
    """提取工期"""
    patterns = [
        r'工期[：:]\s*(\d+\s*日历天)',
        r'工期[：:]\s*(\d+\s*天)',
        r'工期[：:]\s*(\d+\s*个?月)',
        r'工期[：:]\s*([^\n，。；;]{2,20})',
        r'交货期[：:]\s*([^\n，。；;]{2,20})',
        r'(\d+)\s*日历天',
        r'(\d+)\s*个?月',
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1).strip()
    return ''


def _extract_budget(text):
    """提取预算/投资金额"""
    patterns = [
        r'最高投标限价[：:]\s*(\d+(?:\.\d+)?\s*[万亿]?元)',
        r'最高限价[：:]\s*(\d+(?:\.\d+)?\s*[万亿]?元)',
        r'招标控制价[：:]\s*(\d+(?:\.\d+)?\s*[万亿]?元)',
        r'预算金额[：:]\s*(\d+(?:\.\d+)?\s*[万亿]?元)',
        r'投资概算[：:]\s*(\d+(?:\.\d+)?\s*[万亿]?元)',
        r'总投资[：:]\s*(\d+(?:\.\d+)?\s*[万亿]?元)',
        r'合同金额[：:]\s*(\d+(?:\.\d+)?\s*[万亿]?元)',
        r'(\d+(?:\.\d+)?\s*万元)',
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            budget = match.group(1).strip()
            if len(budget) <= 30:
                return budget
    return ''


# ====================== 海况参数提取（增强版） ======================

def _extract_number_near(text, keywords, context_chars=30):
    """
    在关键词附近搜索数值，支持:
    - 范围值: 12~16, 12-16, 12~16米
    - 前缀符号: ≤5, ≥3, <5, >3
    - 纯数值: 15米, 1.5米/秒
    - 数值在关键词前或后: "≤5米的风浪" 或 "波高3m"
    返回匹配到的原始字符串
    """
    for kw in keywords:
        idx = text.find(kw)
        while idx != -1:
            # 取关键词前后的上下文
            ctx_start = max(0, idx - context_chars)
            before_text = text[ctx_start:idx]
            after_text = text[idx + len(kw):idx + len(kw) + context_chars]
            full_context = before_text + kw + after_text

            # 数值模式：可选前缀(≤≥<>=~) + 数字(可带小数) + 可选范围(~-) + 数字
            num_pattern = r'([≤≥<>=～]?)\s*(\d+\.?\d*)\s*[~～\-至到]+\s*(\d+\.?\d*)|' \
                          r'([≤≥<>=～]?)\s*(\d+\.?\d*)'

            # 先在关键词后面找
            match = re.search(num_pattern, after_text)
            if match:
                if match.group(2) and match.group(3):
                    # 范围值
                    prefix = match.group(1) or ''
                    return f'{prefix}{match.group(2)}~{match.group(3)}'
                elif match.group(5):
                    return f'{match.group(4) or ""}{match.group(5)}'

            # 再在关键词前面找（如 "≤5米的风浪" 中5在风浪前面）
            # 反向搜索：找最靠近关键词的数值
            before_matches = list(re.finditer(num_pattern, before_text))
            if before_matches:
                last_match = before_matches[-1]  # 最靠近关键词的
                if last_match.group(2) and last_match.group(3):
                    prefix = last_match.group(1) or ''
                    return f'{prefix}{last_match.group(2)}~{last_match.group(3)}'
                elif last_match.group(5):
                    return f'{last_match.group(4) or ""}{last_match.group(5)}'

            idx = text.find(kw, idx + 1)
    return None


def _extract_sea_conditions(text):
    """提取海况参数（优化版，增强上下文匹配和准确率）"""
    sea = {}

    # 水深 — 多模式匹配，优先完整范围
    # 模式: "水深12~16米" "设计水深15m" "水深12-16米" "水深为12米"
    for pattern in [
        r'水深[为约]?[：:]?\s*(\d+\.?\d*)\s*[~～\-至到]+\s*(\d+\.?\d*)\s*米',
        r'水深[为约]?[：:]?\s*[≤<]?\s*(\d+\.?\d*)\s*米',
        r'设计水深[：:]?\s*(\d+\.?\d*)\s*[~～\-至到]+\s*(\d+\.?\d*)',
        r'设计水深[：:]?\s*[≤<]?\s*(\d+\.?\d*)',
        r'(?:深度|满潮水深)[：:]?\s*[≤<]?\s*(\d+\.?\d*)\s*米',
    ]:
        m = re.search(pattern, text)
        if m:
            groups = [g for g in m.groups() if g]
            if len(groups) >= 2:
                sea['water_depth'] = f'{groups[0]}~{groups[1]}'
            else:
                sea['water_depth'] = groups[0]
            break
    # 回退到 _extract_number_near
    if 'water_depth' not in sea:
        val = _extract_number_near(text, ['水深', '设计水深', '深度', '满潮水深'])
        if val:
            sea['water_depth'] = val

    # 波高 — 多模式匹配
    # 模式: "波高≤5m" "有效波高3米" "Hs=5" "≤5米的风浪"
    for pattern in [
        r'(?:有效波高|最大波高|设计波高|波高)[为约]?[：:]?\s*[≤<]?\s*(\d+\.?\d*)\s*[~～\-至到]+\s*(\d+\.?\d*)\s*米?',
        r'(?:有效波高|最大波高|设计波高|波高)[为约]?[：:]?\s*[≤<]?\s*(\d+\.?\d*)\s*米?',
        r'(?:有效波高|最大波高|设计波高|波高)[为约]?[：:]?\s*[≤<]?\s*(\d+\.?\d*)\s*(?:m|米)',
        r'[≤<]?\s*(\d+\.?\d*)\s*米?\s*的?\s*(?:风浪|波浪|波高)',
        r'Hs[=＝：:]?\s*(\d+\.?\d*)',
    ]:
        m = re.search(pattern, text, re.IGNORECASE)
        if m:
            groups = [g for g in m.groups() if g]
            prefix = '≤' if '≤' in m.group(0) or '<' in m.group(0) else ''
            if len(groups) >= 2:
                sea['max_wave_height'] = f'{prefix}{groups[0]}~{groups[1]}'
            else:
                sea['max_wave_height'] = f'{prefix}{groups[0]}'
            break

    # 流速 — 多模式匹配
    # 模式: "流速≤1.5m/s" "表层流速0.5米/秒" "最大流速1.5"
    for pattern in [
        r'(?:最大流速|表层流速|设计流速|流速)[为约]?[：:]?\s*[≤<]?\s*(\d+\.?\d*)\s*(?:m/s|米/秒|m／s|米／秒)',
        r'(?:最大流速|表层流速|设计流速|流速)[为约]?[：:]?\s*[≤<]?\s*(\d+\.?\d*)\s*[~～\-至到]+\s*(\d+\.?\d*)',
        r'(?:最大流速|表层流速|设计流速|流速)[为约]?[：:]?\s*[≤<]?\s*(\d+\.?\d*)',
        r'[≤<]?\s*(\d+\.?\d*)\s*(?:m/s|米/秒)\s*(?:的)?\s*流速',
    ]:
        m = re.search(pattern, text, re.IGNORECASE)
        if m:
            groups = [g for g in m.groups() if g]
            prefix = '≤' if '≤' in m.group(0) or '<' in m.group(0) else ''
            if len(groups) >= 2:
                sea['max_flow_speed'] = f'{prefix}{groups[0]}~{groups[1]}'
            else:
                sea['max_flow_speed'] = f'{prefix}{groups[0]}'
            break
    # 回退
    if 'max_flow_speed' not in sea:
        val = _extract_number_near(text, ['流速', '水流', '海流'])
        if val:
            sea['max_flow_speed'] = val

    # 风速/台风 — 优化匹配
    # 台风等级: "抗台风≤12级" "≤12级台风" "台风12级"
    for pattern in [
        r'(?:抗台风|抗风|台风|风力)[≤<]?\s*(\d+)\s*级',
        r'[≤<]?\s*(\d+)\s*级.*?(?:台风|风力|风级)',
        r'台风.*?[≤<]?\s*(\d+)\s*级',
    ]:
        m = re.search(pattern, text)
        if m:
            prefix = '≤' if '≤' in m.group(0) or '<' in m.group(0) else ''
            sea['typhoon_level'] = f'{prefix}{m.group(1)}级'
            break

    # 风速: "设计风速30m/s" "风速≤30"
    for pattern in [
        r'(?:设计风速|最大风速|风速|阵风)[为约]?[：:]?\s*[≤<]?\s*(\d+\.?\d*)\s*(?:m/s|米/秒)',
        r'(?:设计风速|最大风速|风速|阵风)[为约]?[：:]?\s*[≤<]?\s*(\d+\.?\d*)',
    ]:
        m = re.search(pattern, text, re.IGNORECASE)
        if m:
            prefix = '≤' if '≤' in m.group(0) or '<' in m.group(0) else ''
            sea['max_wind_speed'] = f'{prefix}{m.group(1)}'
            break
    if 'max_wind_speed' not in sea:
        val = _extract_number_near(text, ['风速', '阵风'])
        if val:
            sea['max_wind_speed'] = val

    # 气温 — 支持 "-20°C至40°C" "-20~40℃"
    m = re.search(r'(-?\d+)\s*[°℃]\s*[Cc]?\s*[至到\-~～]\s*(\d+)\s*[°℃]', text)
    if m:
        sea['temperature_range'] = f'{m.group(1)}°C~{m.group(2)}°C'
    else:
        m = re.search(r'气温[：: ]*(-?\d+)\s*[°℃]', text)
        if m:
            sea['temperature_range'] = f'{m.group(1)}°C'
    # 也匹配 "-20~40°C" 无前置关键词
    if 'temperature_range' not in sea:
        m = re.search(r'(-?\d+)\s*[~～至到\-]\s*(\d+)\s*[°℃]', text)
        if m:
            sea['temperature_range'] = f'{m.group(1)}°C~{m.group(2)}°C'

    # 海域面积 — "海域面积5000亩" "用海面积300公顷"
    m = re.search(r'(?:海域面积|用海面积|养殖面积)\s*(\d+(?:\.\d+)?)\s*(亩|平方米|㎡|公顷|km2|km²)', text)
    if m:
        sea['sea_area_size'] = f'{m.group(1)}{m.group(2)}'

    # 离岸距离 — "离岸约20公里" "距岸15海里"
    m = re.search(r'(?:离岸|距岸|离岸距离)[约]?\s*(\d+(?:\.\d+)?)\s*(公里|千米|km|海里|nm|nmile)', text, re.IGNORECASE)
    if m:
        sea['offshore_distance'] = f'{m.group(1)}{m.group(2)}'

    # 潮差/水位差
    m = re.search(r'潮差[：: ]*(\d+\.?\d*)\s*米?', text)
    if m:
        sea['water_level_diff'] = m.group(1)
    m = re.search(r'水位差[：: ]*(\d+\.?\d*)', text)
    if m:
        sea['water_level_diff'] = m.group(1)

    # 底质
    m = re.search(r'底质[：:]\s*([^\n，。；;]{2,20})', text)
    if m:
        sea['seabed_type'] = m.group(1).strip()
    # 也匹配 "海底底质为"
    if 'seabed_type' not in sea:
        m = re.search(r'海底(?:底质|地质)[为是]?\s*([^\n，。；;]{2,20})', text)
        if m:
            sea['seabed_type'] = m.group(1).strip()

    return sea


# ====================== 网箱参数提取（增强版） ======================

def _extract_cage_params(text):
    """提取网箱参数，区分圆形和方形"""
    cage = {}

    # 管径 — 优先匹配 DNxxx 格式 (DN200, DN 315, dn400 等)
    # 模式1: "管径DN250" "主管径DN315" "主浮管DN400" 等
    m = re.search(r'(?:管径|主管径|主浮管|浮管|浮筒管材?)[：: ]*(DN\s*\d+(?:\s*[×x]\s*DN\s*\d+)?)', text, re.IGNORECASE)
    if m:
        cage['pipe_diameter'] = re.sub(r'\s+', '', m.group(1).upper())
    else:
        # 模式2: 单独出现的 DN+数字 (如 "DN250" "DN 315")
        m = re.search(r'\b(DN\s*\d+)\b', text, re.IGNORECASE)
        if m:
            cage['pipe_diameter'] = re.sub(r'\s+', '', m.group(1).upper())
        else:
            # 模式3: 管径后面跟纯数字+mm (如 "管径315mm")
            m = re.search(r'(?:管径|主管径|主浮管)[：: ]*(\d{2,4})\s*(?:mm|毫米|MM)', text, re.IGNORECASE)
            if m:
                cage['pipe_diameter'] = f'DN{m.group(1)}'
            else:
                # 模式4: "浮管外径为400mm" "浮管外径400mm" "外径为400mm" "管外径400mm"
                m = re.search(r'(?:浮管|主浮管|主管|外径|管外径)[为是]?\s*(\d{2,4})\s*(?:mm|毫米|MM)', text, re.IGNORECASE)
                if m:
                    cage['pipe_diameter'] = f'DN{m.group(1)}'

    # 周长 — "周长60米" "周长60m"
    m = re.search(r'周长\s*(\d+\.?\d*)\s*米?', text)
    if m:
        cage['perimeter'] = float(m.group(1))

    # 直径
    m = re.search(r'直径\s*(\d+\.?\d*)\s*米?', text)
    if m:
        cage['diameter'] = float(m.group(1))

    # 边长（方形网箱）— "边长10米" "边长10m"
    m = re.search(r'边长\s*(\d+\.?\d*)\s*米?', text)
    if m:
        cage['side_length'] = float(m.group(1))

    # 网箱数量 — "64座" "20座...44座" "10口"
    # 优先匹配 "投放XX座"
    m = re.search(r'(?:投放|建设|安装|采购)\s*(\d+)\s*座', text)
    if m:
        cage['cage_count'] = int(m.group(1))
    else:
        m = re.search(r'(\d+)\s*座.*?网箱', text)
        if m:
            cage['cage_count'] = int(m.group(1))
        else:
            m = re.search(r'(\d+)\s*口\s*网?箱?', text)
            if m:
                cage['cage_count'] = int(m.group(1))

    # 网箱类型 — 增强圆形/方形判断
    if re.search(r'圆形|圆型|周长', text) and not re.search(r'方形|方型|矩形', text):
        cage['cage_shape'] = 'circle'
        m = re.search(r'((?:HDPE|PE|橡胶|柔性|休闲型|传统型)?\s*(?:重力式\s*)?(?:深水\s*)?圆形\s*网箱)', text)
        cage['cage_type'] = m.group(1).strip() if m else '圆形网箱'
    elif re.search(r'方形|方型|矩形', text) and not re.search(r'圆形|圆型', text):
        cage['cage_shape'] = 'square'
        m = re.search(r'((?:HDPE|PE|橡胶|柔性)?\s*(?:重力式\s*)?(?:深水\s*)?方形\s*网箱)', text)
        cage['cage_type'] = m.group(1).strip() if m else '方形网箱'
    else:
        # 默认圆形
        m = re.search(r'(HDPE|PE|橡胶|柔性|圆柱形|圆形|方形)\s*(?:重力式)?\s*(?:深水)?\s*网?箱?', text)
        if m:
            cage['cage_type'] = m.group(1)
            cage['cage_shape'] = 'square' if m.group(1) in ('方形',) else 'circle'
        else:
            cage['cage_shape'] = 'circle'

    # 支架间距
    m = re.search(r'支架间距[：: ]*(\d+\.?\d*)\s*米?', text)
    if m:
        cage['bracket_spacing'] = float(m.group(1))

    # 网衣
    m = re.search(r'网衣[：: ]*([^\n，。；;]{3,50})', text)
    if m:
        cage['net_demand'] = m.group(1).strip()

    # 锚固系统
    m = re.search(r'锚固系统[：: ]*([^\n，。；;]{3,80})', text)
    if m:
        cage['mooring_system'] = m.group(1).strip()

    # 锚重
    m = re.search(r'锚(?:重|重量)[：: ]*(\d+)\s*(kg|公斤|吨)', text)
    if m:
        cage['anchor_weight'] = f'{m.group(1)}{m.group(2)}'

    return cage


def _extract_platform_params(text):
    """提取平台参数"""
    platform = {}
    m = re.search(r'面积[：: ]*(\d+\.?\d*)\s*(?:m2|㎡|平方米|平方)', text)
    if m:
        platform['platform_area'] = float(m.group(1))
    m = re.search(r'尺寸[：: ]*([^\n，。；;]{3,30})', text)
    if m:
        platform['size_spec'] = m.group(1).strip()
    return platform


def _extract_breakwater_params(text):
    """提取防波堤参数"""
    breakwater = {}
    m = re.search(r'减波率[：: ]*(\d+%?)', text)
    if m:
        breakwater['wave_reduction_rate'] = m.group(1)
    return breakwater


# ====================== 货物需求提取 ======================

def _extract_goods_requirements(text):
    """
    提取招标文件中的货物需求/采购清单/技术规格部分
    返回: list[dict] 货物需求列表
    """
    goods = []

    # 货物需求关键词
    section_keywords = [
        '货物需求', '采购清单', '货物一览表', '技术规格',
        '规格要求', '设备清单', '材料清单', '货物参数',
        '供货要求', '招标货物', '采购需求', '货物标准',
        '技术要求', '技术参数', '招标范围'
    ]

    # 找到货物需求相关章节
    sections = _find_sections(text, section_keywords)

    for section_title, section_text in sections:
        items = _parse_goods_items(section_text)
        for item in items:
            item['source_section'] = section_title
            goods.append(item)

    # 如果没找到明确章节，尝试从全文搜索
    if not goods:
        goods = _parse_table_style_goods(text)

    # 额外：从"招标范围"段落提取网箱信息
    if not any('网箱' in g.get('name', '') for g in goods):
        cage_goods = _extract_cage_goods_from_text(text)
        goods.extend(cage_goods)

    # 去重
    seen = set()
    unique_goods = []
    for g in goods:
        name = g.get('name', '').strip()
        if name and name not in seen:
            seen.add(name)
            unique_goods.append(g)

    return unique_goods


def _extract_cage_goods_from_text(text):
    """从全文提取网箱类货物信息"""
    goods = []
    # 匹配 "20座周长60米休闲型重力式深水圆形网箱"
    patterns = [
        r'(\d+)\s*座.*?周长\s*(\d+)\s*米.*?((?:休闲型|传统型)?(?:重力式)?(?:深水)?(?:圆形|方形)?网箱)',
        r'(\d+)\s*座.*?((?:休闲型|传统型)?(?:重力式)?(?:深水)?(?:圆形|方形)?网箱).*?周长\s*(\d+)\s*米',
        r'(\d+)\s*(?:座|口|个)\s*(HDPE\s*(?:重力式\s*)?(?:圆形\s*)?网箱)',
    ]
    for pattern in patterns:
        for m in re.finditer(pattern, text):
            if len(m.groups()) >= 2:
                if m.group(2).isdigit():
                    # 第一个模式
                    goods.append({
                        'name': m.group(3) or '网箱',
                        'spec': f'周长{m.group(2)}米',
                        'quantity': m.group(1),
                        'requirements': ''
                    })
                else:
                    goods.append({
                        'name': m.group(2),
                        'spec': '',
                        'quantity': m.group(1),
                        'requirements': ''
                    })
            break  # 只取第一个匹配

    # 锚固系统
    if re.search(r'锚固系统|锚固|系泊', text):
        goods.append({
            'name': '锚固系统',
            'spec': '',
            'quantity': '',
            'requirements': '圆形网箱锚固系统，具体详见技术要求'
        })

    return goods


def _find_sections(text, keywords):
    """在文本中查找包含关键词的章节"""
    sections = []
    lines = text.split('\n')

    for i, line in enumerate(lines):
        line_stripped = line.strip()
        if not line_stripped:
            continue

        for kw in keywords:
            if kw in line_stripped and len(line_stripped) < 30:
                section_lines = [line_stripped]
                for j in range(i + 1, min(i + 100, len(lines))):
                    next_line = lines[j].strip()
                    if re.match(r'^第[一二三四五六七八九十\d]+章', next_line):
                        break
                    if re.match(r'^[一二三四五六七八九十\d]+[.、．]', next_line) and len(next_line) < 30:
                        break
                    if next_line:
                        section_lines.append(next_line)
                if len(section_lines) > 1:
                    sections.append((kw, '\n'.join(section_lines)))
                break

    return sections


def _parse_goods_items(text):
    """从章节文本中解析货物条目"""
    items = []
    lines = text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if '\t' in line:
            cells = [c.strip() for c in line.split('\t') if c.strip()]
            if len(cells) >= 2:
                item = _parse_goods_row(cells)
                if item:
                    items.append(item)
                continue

        cells = line.split()
        if len(cells) >= 3:
            item = _parse_goods_row(cells)
            if item:
                items.append(item)
                continue

        m = re.match(r'(.+?)[：:]\s*(.+)', line)
        if m:
            name = m.group(1).strip()
            spec = m.group(2).strip()
            if _is_likely_goods_name(name):
                items.append({'name': name, 'spec': spec, 'quantity': '', 'requirements': ''})

    return items


def _parse_goods_row(cells):
    """从一行单元格解析货物信息"""
    start = 0
    if cells[0] and re.match(r'^\d+$', cells[0]):
        start = 1

    if start >= len(cells):
        return None

    name = cells[start].strip()
    if not _is_likely_goods_name(name):
        return None

    spec = ''
    quantity = ''
    requirements = ''

    if start + 1 < len(cells):
        spec = cells[start + 1].strip()
    if start + 2 < len(cells):
        qty_text = cells[start + 2].strip()
        if re.match(r'^\d+\.?\d*$', qty_text):
            quantity = qty_text
    if start + 3 < len(cells):
        qty_text = cells[start + 3].strip()
        if re.match(r'^\d+\.?\d*$', qty_text) and not quantity:
            quantity = qty_text
    if start + 4 < len(cells):
        requirements = cells[start + 4].strip()
    elif start + 3 < len(cells) and not re.match(r'^\d+\.?\d*$', cells[start + 3]):
        requirements = cells[start + 3].strip()

    return {'name': name, 'spec': spec, 'quantity': quantity, 'requirements': requirements}


def _is_likely_goods_name(text):
    """判断文本是否像货物名称"""
    if not text or len(text) < 2 or len(text) > 50:
        return False
    if re.match(r'^\d+$', text):
        return False
    if re.match(r'^[一二三四五六七八九十]+[.、]', text):
        return False
    exclude_words = ['序号', '名称', '规格', '数量', '单位', '备注', '合计', '总计',
                     '条款号', '评分因素', '评分标准']
    if text in exclude_words:
        return False
    return True


def _parse_table_style_goods(text):
    """从全文搜索表格样式的货物清单"""
    goods = []
    lines = text.split('\n')
    goods_keywords = ['网箱', '管材', '管件', '网衣', '锚', '绳', '浮筒', '平台',
                      '系泊', '连接件', '钢材', '混凝土', '网具', '浮架', '锚固']

    for line in lines:
        line = line.strip()
        if not line or '\t' not in line:
            continue
        cells = [c.strip() for c in line.split('\t') if c.strip()]
        if len(cells) < 2:
            continue
        for cell in cells:
            if any(kw in cell for kw in goods_keywords):
                item = _parse_goods_row(cells)
                if item:
                    goods.append(item)
                break
    return goods


def _generate_goods_summary(goods_list):
    """生成货物需求总结"""
    if not goods_list:
        return ''
    parts = [f'共识别到 {len(goods_list)} 项货物需求：']
    for i, g in enumerate(goods_list[:10], 1):
        line = f"{i}. {g.get('name', '')}"
        if g.get('spec'):
            line += f"（规格: {g['spec']}）"
        if g.get('quantity'):
            line += f" 数量: {g['quantity']}"
        parts.append(line)
    if len(goods_list) > 10:
        parts.append(f'... 等共 {len(goods_list)} 项')
    return '\n'.join(parts)


def _generate_summary(text, result):
    """生成项目概述（优化版，优先提取项目概况章节）"""

    # 1. 尝试从文本中提取"项目概况""工程概况""项目简介"等章节
    overview = _extract_overview_section(text)

    parts = []

    # 如果提取到了概况章节，作为概述开头
    if overview:
        parts.append(overview)

    # 2. 结构化信息补充
    if result.get('project_scale'):
        parts.append(f"项目规模：{result['project_scale']}")
    if result.get('construction_period'):
        parts.append(f"建设工期：{result['construction_period']}")
    if result.get('budget'):
        parts.append(f"投资预算：{result['budget']}")
    if result.get('tenderer'):
        parts.append(f"招标人：{result['tenderer']}")

    # 海况参数汇总
    sea = result.get('sea_conditions', {})
    sea_parts = []
    if sea.get('water_depth'):
        sea_parts.append(f"水深{sea['water_depth']}米")
    if sea.get('max_wave_height'):
        sea_parts.append(f"波高{sea['max_wave_height']}米")
    if sea.get('max_flow_speed'):
        sea_parts.append(f"流速{sea['max_flow_speed']}m/s")
    if sea.get('typhoon_level'):
        sea_parts.append(f"抗台风{sea['typhoon_level']}")
    if sea.get('max_wind_speed'):
        sea_parts.append(f"风速{sea['max_wind_speed']}m/s")
    if sea_parts:
        parts.append("海况条件：" + "，".join(sea_parts))

    # 网箱参数汇总 — 区分圆形和方形
    cage_circle = result.get('cage_params_circle', {}) or result.get('cage_params', {})
    cage_square = result.get('cage_params_square', {})

    # 圆形网箱
    if cage_circle and (cage_circle.get('cage_count') or cage_circle.get('pipe_diameter')):
        cage_parts = []
        if cage_circle.get('cage_count'):
            cage_parts.append(f"{cage_circle['cage_count']}座")
        if cage_circle.get('perimeter'):
            cage_parts.append(f"周长{cage_circle['perimeter']}米")
        if cage_circle.get('pipe_diameter'):
            cage_parts.append(f"管径{cage_circle['pipe_diameter']}")
        if cage_circle.get('cage_type'):
            cage_parts.append(cage_circle['cage_type'])
        if cage_parts:
            parts.append("圆形网箱：" + "，".join(cage_parts))

    # 方形网箱
    if cage_square and (cage_square.get('cage_count') or cage_square.get('pipe_diameter')):
        cage_parts = []
        if cage_square.get('cage_count'):
            cage_parts.append(f"{cage_square['cage_count']}座")
        if cage_square.get('side_length'):
            cage_parts.append(f"边长{cage_square['side_length']}米")
        if cage_square.get('pipe_diameter'):
            cage_parts.append(f"管径{cage_square['pipe_diameter']}")
        if cage_parts:
            parts.append("方形网箱：" + "，".join(cage_parts))

    return '\n'.join(parts) if parts else ''


def _extract_overview_section(text):
    """
    从招标文件中提取项目概况/项目简介/工程概况等章节内容
    返回简洁的项目概况文本
    """
    # 查找章节标题
    section_titles = [
        '项目概况', '工程概况', '项目简介', '项目背景',
        '工程简介', '项目背景及概况', '建设背景',
    ]

    lines = text.split('\n')
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        for title in section_titles:
            if title in line_stripped and len(line_stripped) < 20:
                # 提取后续段落（最多500字符）
                overview_lines = []
                char_count = 0
                for j in range(i + 1, min(i + 50, len(lines))):
                    next_line = lines[j].strip()
                    # 遇到下一个章节标题则停止
                    if re.match(r'^第[一二三四五六七八九十\d]+章', next_line):
                        break
                    if re.match(r'^[一二三四五六七八九十\d]+[.、．]\s', next_line) and len(next_line) < 30:
                        break
                    if next_line:
                        overview_lines.append(next_line)
                        char_count += len(next_line)
                        if char_count >= 500:
                            break
                if overview_lines:
                    overview_text = ''.join(overview_lines[:5])  # 取前5行拼接
                    # 清理多余空白
                    overview_text = re.sub(r'\s+', ' ', overview_text).strip()
                    if len(overview_text) > 20:
                        return overview_text[:500]
    return ''


def get_supported_formats():
    """返回支持的招标文件格式"""
    return [
        {'ext': '.pdf', 'name': 'PDF文档', 'desc': 'PDF格式（推荐）'},
        {'ext': '.docx', 'name': 'Word文档', 'desc': 'Word 2007+格式'},
        {'ext': '.txt', 'name': '文本文件', 'desc': '纯文本格式'},
        {'ext': '.doc', 'name': 'Word 97-2003', 'desc': '旧版Word格式（兼容性有限）'},
    ]
